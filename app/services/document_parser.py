"""Modular document parsing service for PDF and DOCX formats."""

import io
from typing import Tuple, Optional
import fitz  # PyMuPDF
import docx  # python-docx

from app.schemas.resume import ResumeParseResponse
from app.services.document_validator import validate_resume_file, DocumentValidationError


class DocumentParsingError(Exception):
    """Custom exception raised when text extraction fails or file is corrupted."""
    pass


def extract_pdf_text(content: bytes) -> Tuple[str, int]:
    """
    Extracts text and page count from PDF content bytes using PyMuPDF.

    Args:
        content: Raw bytes of PDF file.

    Returns:
        Tuple of (extracted_text, page_count)

    Raises:
        DocumentParsingError: If PDF is invalid, encrypted, or corrupted.
    """
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        raise DocumentParsingError(f"Failed to open PDF document: {str(e)}") from e

    try:
        if doc.is_encrypted:
            raise DocumentParsingError("PDF is password protected or encrypted.")

        page_count = doc.page_count
        pages_text = []

        for page_num in range(page_count):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            if text:
                pages_text.append(text.strip())

        extracted_text = "\n\n".join(filter(None, pages_text)).strip()
        return extracted_text, page_count
    except DocumentParsingError:
        raise
    except Exception as e:
        raise DocumentParsingError(f"Error extracting text from PDF: {str(e)}") from e
    finally:
        doc.close()


def extract_docx_text(content: bytes) -> Tuple[str, Optional[int]]:
    """
    Extracts text from DOCX content bytes using python-docx.

    Args:
        content: Raw bytes of DOCX file.

    Returns:
        Tuple of (extracted_text, None)

    Raises:
        DocumentParsingError: If DOCX is invalid or corrupted.
    """
    try:
        docx_stream = io.BytesIO(content)
        doc = docx.Document(docx_stream)
    except Exception as e:
        raise DocumentParsingError(f"Failed to open DOCX document: {str(e)}") from e

    try:
        paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

        # Extract text from tables if present
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    table_text.append(" | ".join(row_cells))

        all_text = paragraphs_text + table_text
        extracted_text = "\n".join(all_text).strip()
        return extracted_text, None
    except Exception as e:
        raise DocumentParsingError(f"Error extracting text from DOCX: {str(e)}") from e


def parse_resume_document(filename: str, content: bytes) -> ResumeParseResponse:
    """
    Parses a resume document (PDF or DOCX) and returns structured metadata and text.

    Args:
        filename: Name of the uploaded file.
        content: Raw byte content.

    Returns:
        ResumeParseResponse model.

    Raises:
        DocumentValidationError: For invalid extensions or empty files.
        DocumentParsingError: For corrupted files or extraction failures.
    """
    clean_filename, file_type = validate_resume_file(filename, content)

    if file_type == "pdf":
        text, page_count = extract_pdf_text(content)
    elif file_type == "docx":
        text, page_count = extract_docx_text(content)
    else:
        raise DocumentValidationError(f"Unsupported file type: {file_type}")

    return ResumeParseResponse(
        filename=clean_filename,
        file_type=file_type,
        extracted_text=text,
        character_count=len(text),
        page_count=page_count
    )
