"""Pytest configuration and shared fixtures for test suite isolation."""

import os
import io
import fitz
import docx
import pytest
from typing import Generator
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
from fastapi.testclient import TestClient

from app.main import app
from app.db.database import Base, get_db
from app.db.models import ResumeAnalysis, JobAnalysis, MatchAnalysis

# Determine test database URL independently of development DATABASE_URL
DEFAULT_TEST_DB_URL = "postgresql+psycopg://saijoshi@localhost:5432/resumeai_test"
TEST_DATABASE_URL = os.getenv("TEST_DATABASE_URL", DEFAULT_TEST_DB_URL)

# Safety assertion: Ensure test suite NEVER runs against primary development database 'resumeai'
parsed_db_name = TEST_DATABASE_URL.rstrip("/").split("/")[-1]
assert parsed_db_name != "resumeai", (
    f"SAFETY ERROR: Test suite is configured to target primary development database '{parsed_db_name}'. "
    "Tests MUST use an isolated test database (e.g. 'resumeai_test' or SQLite in-memory)."
)

test_engine_args = {}
if TEST_DATABASE_URL.startswith("sqlite"):
    test_engine_args["connect_args"] = {"check_same_thread": False}

test_engine = create_engine(TEST_DATABASE_URL, pool_pre_ping=True, **test_engine_args)
TestSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


@pytest.fixture(scope="session", autouse=True)
def setup_test_database():
    """Session-level setup and teardown for the test database schema."""
    Base.metadata.create_all(bind=test_engine)
    yield
    Base.metadata.drop_all(bind=test_engine)


@pytest.fixture(autouse=True)
def db_isolation() -> Generator[Session, None, None]:
    """
    Per-test isolation fixture automatically active for all tests.
    Overrides FastAPI get_db dependency and cleans test tables before and after each test function.
    """
    session = TestSessionLocal()

    # Truncate/delete test data before test execution
    try:
        session.query(MatchAnalysis).delete()
        session.query(ResumeAnalysis).delete()
        session.query(JobAnalysis).delete()
        session.commit()
    except Exception:
        session.rollback()

    def override_get_db():
        try:
            yield session
        finally:
            pass

    app.dependency_overrides[get_db] = override_get_db

    try:
        yield session
    finally:
        session.rollback()
        # Clean up after test execution
        try:
            session.query(MatchAnalysis).delete()
            session.query(ResumeAnalysis).delete()
            session.query(JobAnalysis).delete()
            session.commit()
        except Exception:
            session.rollback()
        session.close()
        app.dependency_overrides.clear()


@pytest.fixture
def client(db_isolation: Session) -> TestClient:
    """Provides a FastAPI TestClient instance configured with test database isolation."""
    return TestClient(app)


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Generates valid sample PDF file bytes with text."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        "John Doe\njohn@doe.dev | 555-0199\nSUMMARY\nSoftware Engineer\nEXPERIENCE\nExperience in Python, FastAPI, and Machine Learning.\nEDUCATION\nBS Computer Science\nSKILLS\nPython, FastAPI, SQL"
    )
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Generates valid sample DOCX file bytes with text."""
    doc = docx.Document()
    doc.add_heading("Jane Smith", level=0)
    doc.add_paragraph("jane@smith.dev | 555-0188")
    doc.add_paragraph("SUMMARY\nSenior Data Scientist")
    doc.add_paragraph("EXPERIENCE\nExpertise in NLP, Python, and PyTorch.")
    doc.add_paragraph("EDUCATION\nMS Data Science")
    doc.add_paragraph("SKILLS\nPython, PyTorch, SQL")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
